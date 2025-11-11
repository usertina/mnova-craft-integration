"""
Utilidades de Exportación Completas - CraftRMN Pro v2.3 FINAL
Genera reportes completos con gráficos en PDF, DOCX y CSV.
Soporta: Reportes individuales, Comparación de muestras, Dashboard.
Idiomas: Castellano, English, Euskera, Galego

VERSIÓN FINAL con todas las correcciones:
✅ Traducciones correctas en 4 idiomas
✅ Extracción robusta de concentración y confianza de compuestos PFAS
✅ Sección de Compuestos PFAS Detectados completa
✅ Todas las funciones incluidas (individual, comparison, dashboard)
"""

import io
import json
import base64
from datetime import datetime
from pathlib import Path
from typing import Dict, List, BinaryIO, Optional 
import os

# --- Plotly ---
try:
    import plotly.graph_objects as go
    from plotly.io import to_image
except ImportError:
    print("Error: Plotly no está instalado. `pip install plotly kaleido`")
    go = None

# --- Traducciones ---
try:
    from translation_manager import TranslationManager
except ImportError:
    class TranslationManager:
        def __init__(self, lang='es'):
            self.lang = lang
        
        def get(self, key, default=None):
            return default or f"[{key}]"
        
        def t(self, key, default=None):
            return default or f"[{key}]"
    print("Advertencia: translation_manager no encontrado. Usando traducciones placeholder.")

# --- DOCX ---
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("Advertencia: python-docx no está instalado. Exportación DOCX deshabilitada.")
    DOCX_AVAILABLE = False
    Document = None

# --- PDF (ReportLab) ---
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, BaseDocTemplate, Frame, PageTemplate,
        Table, TableStyle, Paragraph, Spacer, Image, PageBreak, KeepTogether
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    PDF_AVAILABLE = True
except ImportError:
    print("Advertencia: reportlab no está instalado. Exportación PDF deshabilitada.")
    PDF_AVAILABLE = False
    BaseDocTemplate = None

# ============================================================================
# FUNCIÓN AUXILIAR PARA NORMALIZAR UNICODE
# ============================================================================
def normalize_region_text(text: str) -> str:
    """Convierte caracteres Unicode específicos a texto normal."""
    replacements = {
        '\u2082': '2', '\u2083': '3', '\u207b': '-',
        '₂': '2', '₃': '3', '⁻': '-', 'ó': 'o',
    }
    result = str(text)
    for unicode_char, replacement in replacements.items():
        result = result.replace(unicode_char, replacement)
    return result

# ============================================================================
# CLASE PRINCIPAL ReportExporter
# ============================================================================
class ReportExporter:
    """Clase para exportar análisis RMN con gráficos, branding y soporte multiidioma."""

    # ========================================================================
    # EXPORTACIÓN JSON
    # ========================================================================
    @staticmethod
    def export_json(results: Dict, lang: str = 'es') -> BinaryIO:
        """Exporta resultados como JSON."""
        output = io.BytesIO()
        json_str = json.dumps(results, indent=2, ensure_ascii=False)
        output.write(json_str.encode('utf-8'))
        output.seek(0)
        return output

    # ========================================================================
    # UTILIDADES PARA GRÁFICOS
    # ========================================================================
    @staticmethod
    def plotly_to_image(plotly_json: Dict, format='png', width=800, height=500) -> bytes:
        """Convierte un gráfico Plotly (JSON) a imagen."""
        if go is None: return None
        try:
            fig = go.Figure(plotly_json)
            img_bytes = to_image(fig, format=format, width=width, height=height, scale=2)
            return img_bytes
        except Exception as e:
            print(f"Error convirtiendo gráfico Plotly: {e}")
            return None

    @staticmethod
    def base64_to_bytes(base64_str: str) -> bytes:
        """Convierte string base64 a bytes, manejando padding."""
        if not base64_str: return None
        try:
            if isinstance(base64_str, bytes): base64_str = base64_str.decode('utf-8')
            if ',' in base64_str: base64_str = base64_str.split(',')[1]
            missing_padding = len(base64_str) % 4
            if missing_padding: base64_str += '=' * (4 - missing_padding)
            return base64.b64decode(base64_str)
        except Exception as e:
            print(f"Error decodificando base64: {e}")
            return None

    # ========================================================================
    # FUNCIÓN AUXILIAR PARA PIE DE PÁGINA PDF
    # ========================================================================
    @staticmethod
    def _add_pdf_footer(canvas, doc, company_data):
        """Añade pie de página con info de empresa y número de página."""
        if not company_data: company_data = {}
        canvas.saveState()
        styles = getSampleStyleSheet()
        footer_style = ParagraphStyle(
            'FooterStyle', parent=styles['Normal'], fontSize=8,
            alignment=TA_CENTER, textColor=colors.grey
        )
        company_name = company_data.get('name', 'CraftRMN Pro')
        company_address = company_data.get('address', '')
        company_phone = company_data.get('phone', '')
        company_email = company_data.get('email', '')
        parts = [part for part in [company_name, company_address, f"Tel: {company_phone}", company_email] if part]
        footer_text = " | ".join(parts)
        p = Paragraph(footer_text, footer_style)
        p.wrapOn(canvas, doc.width, doc.bottomMargin)
        p.drawOn(canvas, doc.leftMargin, 0.5 * inch)
        page_num_text = f"Página {canvas.getPageNumber()}"
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.grey)
        canvas.drawRightString(doc.width + doc.leftMargin - 0.1*inch, 0.5 * inch, page_num_text)
        canvas.restoreState()

    # ========================================================================
    # FUNCIÓN AUXILIAR PARA PIE DE PÁGINA DOCX
    # ========================================================================
    @staticmethod
    def _add_docx_footer(doc: Document, company_data: Dict): # type: ignore
        """Añade pie de página con info de empresa a un documento DOCX."""
        if not DOCX_AVAILABLE or not doc: return
        if not company_data: company_data = {}
        try:
            section = doc.sections[0]
            footer = section.footer
            if footer.paragraphs: 
                footer_para = footer.paragraphs[0]
                footer_para.clear()
            else: 
                footer_para = footer.add_paragraph()
            company_name = company_data.get('name', 'CraftRMN Pro')
            company_address = company_data.get('address', '')
            company_phone = company_data.get('phone', '')
            company_email = company_data.get('email', '')
            parts = [part for part in [company_name, company_address, f"Tel: {company_phone}", company_email] if part]
            run = footer_para.add_run(" | ".join(parts))
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e: 
            print(f"Error al configurar pie de página DOCX: {e}")

    # ========================================================================
    # FUNCIÓN AUXILIAR PARA LOGO
    # ========================================================================
    @staticmethod
    def _add_logo(story_or_doc, company_data: Dict, format: str):
        """Añade el logo de la empresa al story (PDF) o doc (DOCX)."""
        if not company_data: return
        company_logo_path = company_data.get('logo_path_on_server')
        if company_logo_path and Path(company_logo_path).exists():
            try:
                if format == 'pdf' and PDF_AVAILABLE:
                    logo_img = Image(company_logo_path, width=1.5*inch, height=0.75*inch)
                    logo_img.hAlign = 'LEFT'
                    story_or_doc.append(logo_img)
                    story_or_doc.append(Spacer(1, 0.2*inch))
                elif format == 'docx' and DOCX_AVAILABLE:
                    story_or_doc.add_picture(company_logo_path, width=Inches(1.5))
                    last_paragraph = story_or_doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    story_or_doc.add_paragraph()
            except Exception as e: 
                print(f"Error al añadir logo {format.upper()} {company_logo_path}: {e}")
        else: 
            print(f"Advertencia {format.upper()}: Logo no encontrado en: {company_logo_path}")

    # ========================================================================
    # ✅ FUNCIÓN: Extraer Valores de Forma Robusta
    # ========================================================================
    @staticmethod
    def _extract_value_safely(results: Dict, analysis: Dict, keys: List[str], default=0):
        """
        Extrae un valor de múltiples ubicaciones posibles en el diccionario.
        Intenta en orden: analysis[key], results[key], results['parameters'][key]
        """
        for key in keys:
            # Intentar en analysis
            val = analysis.get(key)
            if val is not None:
                return val
            
            # Intentar en results
            val = results.get(key)
            if val is not None:
                return val
            
            # Intentar en results['parameters']
            val = results.get('parameters', {}).get(key)
            if val is not None:
                return val
        
        return default

    # ========================================================================
    # EXPORTACIÓN PDF INDIVIDUAL - VERSIÓN FINAL CORREGIDA
    # ========================================================================
    @staticmethod
    def export_pdf(results: Dict, company_data: Dict, chart_image: bytes = None, lang: str = 'es') -> BinaryIO:
        """Exporta análisis individual como PDF con gráfico, datos de empresa y COMPUESTOS PFAS."""
        if not PDF_AVAILABLE: raise ImportError("ReportLab no está instalado.")
        t = TranslationManager(lang)
        output = io.BytesIO()

        doc = BaseDocTemplate(output, pagesize=A4, topMargin=0.75*inch, bottomMargin=1.0*inch, 
                              leftMargin=0.75*inch, rightMargin=0.75*inch)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
        template = PageTemplate(id='main', frames=[frame], 
                                onPage=lambda canvas, doc: ReportExporter._add_pdf_footer(canvas, doc, company_data))
        doc.addPageTemplates([template])
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, 
                                     textColor=colors.HexColor('#2c3e50'), spaceAfter=30, alignment=TA_CENTER)
        subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], alignment=TA_CENTER)
        info_style = ParagraphStyle('Info', parent=styles['Normal'], alignment=TA_CENTER, fontSize=12)

        # PORTADA
        story.append(Paragraph(t.t('report.title', "NMR Analysis Report"), title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(t.t('report.subtitle', "PFAS Detection and Quantification"), subtitle_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"<b>{t.t('report.sample', 'Sample')}:</b> {results.get('filename', 'N/A')}", info_style))
        story.append(Paragraph(f"<b>{t.t('report.date', 'Date')}:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", info_style))
        story.append(PageBreak())

        # LOGO
        ReportExporter._add_logo(story, company_data, 'pdf')

        # GRÁFICO
        story.append(Paragraph(f"1. {t.t('report.spectrum', 'NMR Spectrum')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        if chart_image:
            try: 
                img = Image(io.BytesIO(chart_image), width=6*inch, height=3.5*inch)
                img.hAlign = 'CENTER'
                story.append(KeepTogether(img))
            except Exception as e: 
                print(f"Error al añadir imagen del gráfico PDF: {e}")
                story.append(Paragraph(f"[{t.t('errors.chartError', 'Error displaying chart')}]", styles['Italic']))
        else: 
            story.append(Paragraph(f"[{t.t('errors.noChartData', 'Chart not available')}]", styles['Italic']))
        story.append(Spacer(1, 0.3*inch))

        # RESUMEN EJECUTIVO
        story.append(Paragraph(f"2. {t.t('report.executive_summary', 'Executive Summary')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        analysis = results.get("analysis", {})
        
        fluor_percentage = ReportExporter._extract_value_safely(
            results, analysis, ['fluor_percentage', 'fluor'], 0
        )
        
        pfas_percentage = ReportExporter._extract_value_safely(
            results, analysis, ['pifas_percentage', 'pfas_percentage', 'pfas'], 0
        )
        
        concentration_val = ReportExporter._extract_value_safely(
            results, analysis, ['pifas_concentration', 'pfas_concentration', 'concentration'], 0
        )
        
        summary_data = [
            [t.t('report.parameter', 'Parameter'), 
             t.t('report.value', 'Value'), 
             t.t('report.unit', 'Unit')],
            [t.t('results.fluor', 'Total Fluorine'), 
             f"{fluor_percentage:.2f}", 
             t.t('units.percentage', '%')],
            [t.t('results.pfas', 'PFAS'), 
             f"{pfas_percentage:.2f}", 
             t.t('units.percent_fluor', '% Fluorine')],
            [t.t('results.concentration', 'PFAS Concentration'), 
             f"{concentration_val:.4f}", 
             t.t('units.millimolar', 'mM')],
        ]

        summary_table = Table(summary_data, colWidths=[3*inch, 1.5*inch, 1*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        quality_score = results.get('quality_score', 0)
        quality_text = f"<b>{t.t('report.quality_score', 'Quality Score')}:</b> {quality_score:.1f}{t.t('units.over_ten', '/10')}"
        story.append(Paragraph(quality_text, styles['Normal']))
        story.append(PageBreak())

        # INFORMACIÓN RÁPIDA
        story.append(Paragraph(f"3. {t.t('report.quick_info', 'Quick Info')}", styles['Heading1']))
        story.append(Spacer(1, 0.1*inch))
        
        qm = results.get("quality_metrics", {})
        snr_value = ReportExporter._extract_value_safely(results, qm, ['snr', 'signal_to_noise'], 0)
        total_area = ReportExporter._extract_value_safely(results, analysis, ['total_area', 'total_integral'], 0)
        
        quick_info_data = [
            [t.t('report.metric', 'Metric'), t.t('report.value', 'Value')],
            [t.t('peaks.title', 'Detected Peaks'), str(len(results.get('peaks', [])))],
            [t.t('results.total_area', 'Total Integrated Area'), f"{total_area:,.2f}"],
            ['SNR', f"{snr_value:.2f}" if snr_value else 'N/A']
        ]
        
        quick_table = Table(quick_info_data, colWidths=[2.5*inch, 2*inch])
        quick_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9b59b6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(quick_table)
        story.append(PageBreak())

        # ANÁLISIS DETALLADO
        story.append(Paragraph(f"4. {t.t('report.detailed_analysis', 'Detailed Analysis')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"4.1 {t.t('report.chemical_composition', 'Chemical Composition')}", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        
        fluor_area = ReportExporter._extract_value_safely(results, analysis, ['fluor_area'], 0)
        pfas_area = ReportExporter._extract_value_safely(results, analysis, ['pifas_area', 'pfas_area'], 0)
        sample_concentration = ReportExporter._extract_value_safely(
            results, analysis, ['concentration', 'sample_concentration'], 1.0
        )
        
        detailed_data = [
            [t.t('report.parameter', 'Parameter'), 
             t.t('report.value', 'Value'), 
             t.t('report.unit', 'Unit')],
            [t.t('results.total_area', 'Total Integrated Area'), 
             f"{total_area:,.2f}", 
             t.t('units.arbitrary', 'a.u.')],
            [t.t('results.fluor_area', 'Fluorine Area'), 
             f"{fluor_area:,.2f}", 
             t.t('units.arbitrary', 'a.u.')],
            [t.t('results.pfas_area', 'PFAS Area'), 
             f"{pfas_area:,.2f}", 
             t.t('units.arbitrary', 'a.u.')],
            [t.t('results.sample_concentration', 'Sample Concentration'), 
             f"{sample_concentration:.2f}", 
             t.t('units.millimolar', 'mM')],
        ]
        
        detailed_table = Table(detailed_data, colWidths=[3*inch, 1.5*inch, 1*inch])
        detailed_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e74c3c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(detailed_table)
        story.append(Spacer(1, 0.3*inch))
        
        story.append(Paragraph(f"4.2 {t.t('report.detailed_statistics', 'Detailed Statistics')}", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        
        spectrum_data = results.get('spectrum', {})
        ppm_min = spectrum_data.get('ppm_min', 0)
        ppm_max = spectrum_data.get('ppm_max', 0)
        ppm_range_text = f"{ppm_min:.1f} {t.t('report.to', 'to')} {ppm_max:.1f} ppm" if ppm_min != 0 or ppm_max != 0 else "N/A"
        
        stats_data = [
            [t.t('report.parameter', 'Parameter'), t.t('report.value', 'Value')],
            ['SNR', f"{snr_value:.2f}" if snr_value is not None else 'N/A'],
            [t.t('report.limits', 'Limits'), ppm_range_text],
        ]
        
        stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#16a085')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(stats_table)
        story.append(PageBreak())

        # =================================================================
        # === INICIO DE LA SECCIÓN MODIFICADA (PFAS CON IMÁGENES 2D) ===
        # =================================================================
        
        pfas_detection = results.get('pfas_detection', {})
        
        # --- CORRECCIÓN 1: Buscar 'detected_pfas' (de app.py) si 'compounds' está vacío ---
        compounds = pfas_detection.get('compounds', [])
        if not compounds:
            compounds = pfas_detection.get('detected_pfas', [])
        
        story.append(Paragraph(f"5. {t.t('report.detected_compounds', 'Detected PFAS Compounds')}", 
                               styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))

        if compounds:
            # --- CORRECCIÓN 2: Lógica para incluir imágenes 2D ---
            # Ya no usamos una sola tabla. Iteramos y creamos bloques.
            
            # Estilos para el texto del compuesto
            compound_name_style = ParagraphStyle('CompoundName', parent=styles['Heading3'], fontSize=12, spaceAfter=6, textColor=colors.HexColor('#27ae60'))
            compound_detail_style = ParagraphStyle('CompoundDetail', parent=styles['Normal'], fontSize=10, leftIndent=12, spaceAfter=3)

            for compound in compounds:
                compound_block = [] # Usar KeepTogether para este bloque
                
                name = compound.get('name', 'N/A')
                cas = compound.get('cas_number', compound.get('cas', 'N/A'))
                formula = compound.get('formula', compound.get('molecular_formula', 'N/A'))
                
                # Obtener confianza
                confidence_val = (compound.get('confidence_level') or 
                                  compound.get('confidence') or 
                                  compound.get('match_quality'))

                # Formatear confianza (app.py ya multiplica por 100)
                if confidence_val:
                    try:
                        confidence_text = f"{float(confidence_val):.1f}%"
                    except (ValueError, TypeError):
                        confidence_text = str(confidence_val)
                else:
                    confidence_text = t.t('compounds.detected', 'Detected')

                # Añadir texto
                compound_block.append(Paragraph(name, compound_name_style))
                compound_block.append(Paragraph(
                    f"<b>{t.t('compounds.confidence', 'Confidence')}:</b> {confidence_text}", 
                    compound_detail_style
                ))
                compound_block.append(Paragraph(
                    f"<b>{t.t('compounds.formula', 'Formula')}:</b> {formula}", 
                    compound_detail_style
                ))
                compound_block.append(Paragraph(
                    f"<b>{t.t('compounds.cas', 'CAS')}:</b> {cas}", 
                    compound_detail_style
                ))
                
                # Añadir imagen 2D si existe
                image_b64 = compound.get('image_2d')
                if image_b64:
                    try:
                        img_bytes = ReportExporter.base64_to_bytes(image_b64)
                        if img_bytes:
                            compound_block.append(Spacer(1, 0.1*inch))
                            # Ajustar tamaño de imagen si es necesario
                            img = Image(io.BytesIO(img_bytes), width=1.5*inch, height=1.5*inch, kind='bound') 
                            img.hAlign = 'LEFT'
                            img.leftIndent = 12
                            compound_block.append(img)
                    except Exception as img_err:
                        # No añadir nada si la imagen falla
                        print(f"Advertencia: No se pudo renderizar la imagen 2D para {name}: {img_err}")
                
                compound_block.append(Spacer(1, 0.25*inch)) # Espacio después de cada compuesto
                story.append(KeepTogether(compound_block))
            
            # --- FIN CORRECCIÓN 2 ---
        
        else:
            # Mensaje 'No detectado'
            story.append(Paragraph(t.t('compounds.none', 'No specific PFAS compounds were detected.'),
                                   styles['Normal']))
        
        # =================================================================
        # === FIN DE LA SECCIÓN MODIFICADA ===
        # =================================================================
        
        story.append(PageBreak())

        # PICOS DETECTADOS
        story.append(Paragraph(f"6. {t.t('report.detected_peaks', 'Detected Peaks')}", styles['Heading1']))
        peaks = results.get("peaks", [])
        
        if peaks:
            story.append(Spacer(1, 0.2*inch))
            peaks_data = [[
                t.t('peaks.ppm', 'PPM'),
                t.t('peaks.intensity', 'Intensity'),
                t.t('peaks.relative_intensity', 'Rel. Int.'),
                t.t('peaks.width', 'Width (ppm)'),
                t.t('peaks.width_hz', 'Width (Hz)'),
                'SNR',
                t.t('peaks.region', 'Region')
            ]]
            
            for peak in peaks:
                ppm_val = peak.get('ppm') or peak.get('position') or 0
                intensity_val = peak.get('intensity') or peak.get('height') or 0
                rel_intensity_val = peak.get('relative_intensity') or 0
                width_ppm_val = peak.get('width_ppm') or peak.get('width') or 0
                width_hz_val = peak.get('width_hz') or 0
                snr_val = peak.get('snr') or 0
                region_val = normalize_region_text(peak.get('region', 'N/A'))
                
                peaks_data.append([
                    f"{ppm_val:.3f}",
                    f"{intensity_val:,.0f}",
                    f"{rel_intensity_val:.1f}%",
                    f"{width_ppm_val:.3f}",
                    f"{width_hz_val:.1f}",
                    f"{snr_val:.2f}",
                    region_val
                ])

            peaks_table = Table(peaks_data, colWidths=[0.9*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch, 1.4*inch])
            # Ajustar anchos si el contenido se corta (ejemplo anterior tenía 7 columnas)
            peaks_table = Table(peaks_data, colWidths=[0.9*inch, 1.0*inch, 0.8*inch, 0.9*inch, 0.7*inch, 0.7*inch, 1.1*inch])
            
            peaks_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f39c12')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('WORDWRAP', (6, 1), (6, -1), True),
            ]))
            story.append(peaks_table)
        else:
            story.append(Paragraph(t.t('peaks.none', 'No significant peaks detected.'), 
                                   styles['Normal']))

        doc.build(story)
        output.seek(0)
        return output
    # ========================================================================
    # EXPORTACIÓN DOCX INDIVIDUAL - VERSIÓN FINAL CORREGIDA
    # ========================================================================
    @staticmethod
    def export_docx(results: Dict, company_data: Dict, chart_image: bytes = None, lang: str = 'es') -> BinaryIO:
        """Exporta análisis individual como DOCX con gráfico, datos de empresa y COMPUESTOS PFAS."""
        if not DOCX_AVAILABLE: 
            raise ImportError("python-docx no está instalado.")
        
        t = TranslationManager(lang)
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # PIE DE PÁGINA
        ReportExporter._add_docx_footer(doc, company_data)

        # PORTADA
        title = doc.add_heading(t.t('report.title', "NMR Analysis Report"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(t.t('report.subtitle', "PFAS Detection and Quantification"), level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        info = doc.add_paragraph()
        info.add_run(f"{t.t('report.sample', 'Sample')}: ").bold = True
        info.add_run(results.get('filename', 'N/A'))
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_p = doc.add_paragraph()
        date_p.add_run(f"{t.t('report.date', 'Date')}: ").bold = True
        date_p.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()

        # LOGO
        ReportExporter._add_logo(doc, company_data, 'docx')

        # GRÁFICO
        doc.add_heading(f"1. {t.t('report.spectrum', 'NMR Spectrum')}", level=1)
        
        if chart_image:
            try:
                doc.add_picture(io.BytesIO(chart_image), width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error al añadir imagen del gráfico DOCX: {e}")
                doc.add_paragraph(f"[{t.t('errors.chartError', 'Error displaying chart')}]")
        else:
            doc.add_paragraph(f"[{t.t('errors.noChartData', 'Chart not available')}]")
        
        doc.add_paragraph()

        # RESUMEN EJECUTIVO
        doc.add_heading(f"2. {t.t('report.executive_summary', 'Executive Summary')}", level=1)
        
        analysis = results.get("analysis", {})
        
        fluor_percentage = ReportExporter._extract_value_safely(
            results, analysis, ['fluor_percentage', 'fluor'], 0
        )
        
        pfas_percentage = ReportExporter._extract_value_safely(
            results, analysis, ['pifas_percentage', 'pfas_percentage', 'pfas'], 0
        )
        
        concentration_val = ReportExporter._extract_value_safely(
            results, analysis, ['pifas_concentration', 'pfas_concentration', 'concentration'], 0
        )
        
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = t.t('report.parameter', 'Parameter')
        hdr_cells[1].text = t.t('report.value', 'Value')
        hdr_cells[2].text = t.t('report.unit', 'Unit')
        
        data_rows = [
            (t.t('results.fluor', 'Total Fluorine'), f"{fluor_percentage:.2f}", t.t('units.percentage', '%')),
            (t.t('results.pfas', 'PFAS'), f"{pfas_percentage:.2f}", t.t('units.percent_fluor', '% Fluorine')),
            (t.t('results.concentration', 'PFAS Concentration'), f"{concentration_val:.4f}", t.t('units.millimolar', 'mM')),
        ]
        
        for idx, (param, value, unit) in enumerate(data_rows, 1):
            if idx < len(table.rows):
                row_cells = table.rows[idx].cells
                row_cells[0].text = param
                row_cells[1].text = value
                row_cells[2].text = unit
        
        doc.add_paragraph()
        
        quality_score = results.get('quality_score', 0)
        quality_p = doc.add_paragraph()
        quality_p.add_run(f"{t.t('report.quality_score', 'Quality Score')}: ").bold = True
        quality_p.add_run(f"{quality_score:.1f}{t.t('units.over_ten', '/10')}")
        
        doc.add_page_break()

        # INFORMACIÓN RÁPIDA
        doc.add_heading(f"3. {t.t('report.quick_info', 'Quick Info')}", level=1)
        
        qm = results.get("quality_metrics", {})
        snr_value = ReportExporter._extract_value_safely(results, qm, ['snr', 'signal_to_noise'], 0)
        total_area = ReportExporter._extract_value_safely(results, analysis, ['total_area', 'total_integral'], 0)
        
        quick_table = doc.add_table(rows=4, cols=2)
        quick_table.style = 'Light Grid Accent 1'
        
        quick_hdr = quick_table.rows[0].cells
        quick_hdr[0].text = t.t('report.metric', 'Metric')
        quick_hdr[1].text = t.t('report.value', 'Value')
        
        quick_rows = [
            (t.t('peaks.title', 'Detected Peaks'), str(len(results.get('peaks', [])))),
            (t.t('results.total_area', 'Total Integrated Area'), f"{total_area:,.2f}"),
            ('SNR', f"{snr_value:.2f}" if snr_value else 'N/A')
        ]
        
        for idx, (metric, value) in enumerate(quick_rows, 1):
            if idx < len(quick_table.rows):
                row_cells = quick_table.rows[idx].cells
                row_cells[0].text = metric
                row_cells[1].text = value
        
        doc.add_page_break()

        # ANÁLISIS DETALLADO
        doc.add_heading(f"4. {t.t('report.detailed_analysis', 'Detailed Analysis')}", level=1)
        doc.add_heading(f"4.1 {t.t('report.chemical_composition', 'Chemical Composition')}", level=2)
        
        fluor_area = ReportExporter._extract_value_safely(results, analysis, ['fluor_area'], 0)
        pfas_area = ReportExporter._extract_value_safely(results, analysis, ['pifas_area', 'pfas_area'], 0)
        sample_concentration = ReportExporter._extract_value_safely(
            results, analysis, ['concentration', 'sample_concentration'], 1.0
        )
        
        detailed_table = doc.add_table(rows=5, cols=3)
        detailed_table.style = 'Light Grid Accent 1'
        
        hdr_cells_det = detailed_table.rows[0].cells
        hdr_cells_det[0].text = t.t('report.parameter', 'Parameter')
        hdr_cells_det[1].text = t.t('report.value', 'Value')
        hdr_cells_det[2].text = t.t('report.unit', 'Unit')
        
        detailed_rows = [
            (t.t('results.total_area', 'Total Integrated Area'), f"{total_area:,.2f}", t.t('units.arbitrary', 'a.u.')),
            (t.t('results.fluor_area', 'Fluorine Area'), f"{fluor_area:,.2f}", t.t('units.arbitrary', 'a.u.')),
            (t.t('results.pfas_area', 'PFAS Area'), f"{pfas_area:,.2f}", t.t('units.arbitrary', 'a.u.')),
            (t.t('results.sample_concentration', 'Sample Concentration'), f"{sample_concentration:.2f}", t.t('units.millimolar', 'mM')),
        ]
        
        for idx, (param, value, unit) in enumerate(detailed_rows, 1):
            if idx < len(detailed_table.rows):
                row_cells = detailed_table.rows[idx].cells
                row_cells[0].text = param
                row_cells[1].text = value
                row_cells[2].text = unit
        
        doc.add_paragraph()
        
        # ESTADÍSTICAS DETALLADAS
        doc.add_heading(f"4.2 {t.t('report.detailed_statistics', 'Detailed Statistics')}", level=2)
        
        spectrum_data = results.get('spectrum', {})
        ppm_min = spectrum_data.get('ppm_min', 0)
        ppm_max = spectrum_data.get('ppm_max', 0)
        ppm_range_text = f"{ppm_min:.1f} {t.t('report.to', 'to')} {ppm_max:.1f} ppm" if ppm_min != 0 or ppm_max != 0 else "N/A"
        
        stats_table = doc.add_table(rows=3, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        stats_hdr = stats_table.rows[0].cells
        stats_hdr[0].text = t.t('report.parameter', 'Parameter')
        stats_hdr[1].text = t.t('report.value', 'Value')
        
        stats_rows = [
            ('SNR', f"{snr_value:.2f}" if snr_value else 'N/A'),
            (t.t('report.limits', 'Limits'), ppm_range_text),
        ]
        
        for idx, (param, value) in enumerate(stats_rows, 1):
            if idx < len(stats_table.rows):
                row_cells = stats_table.rows[idx].cells
                row_cells[0].text = param
                row_cells[1].text = value
        
        doc.add_page_break()

        # =================================================================
        # === INICIO DE LA SECCIÓN MODIFICADA (PFAS CON IMÁGENES 2D) ===
        # =================================================================
        
        pfas_detection = results.get('pfas_detection', {})
        
        # --- CORRECCIÓN 1: Clave de diccionario ---
        compounds = pfas_detection.get('compounds', [])
        if not compounds:
             compounds = pfas_detection.get('detected_pfas', [])
        
        doc.add_heading(f"5. {t.t('report.detected_compounds', 'Detected PFAS Compounds')}", level=1)
        
        if compounds:
            # --- CORRECCIÓN 2: Lógica para incluir imágenes 2D ---
            # Borrar la creación de la tabla, iterar por compuesto
            
            for compound in compounds:
                name = compound.get('name', 'N/A')
                cas = compound.get('cas_number', compound.get('cas', 'N/A'))
                formula = compound.get('formula', compound.get('molecular_formula', 'N/A'))
                
                confidence_val = (compound.get('confidence_level') or 
                                  compound.get('confidence') or 
                                  compound.get('match_quality'))
                
                if confidence_val:
                    try:
                        confidence_text = f"{float(confidence_val):.1f}%"
                    except (ValueError, TypeError):
                        confidence_text = str(confidence_val)
                else:
                    confidence_text = t.t('compounds.detected', 'Detected')
                
                # Añadir texto
                doc.add_heading(name, level=3)
                
                p = doc.add_paragraph()
                p.add_run(f"{t.t('compounds.confidence', 'Confidence')}: ").bold = True
                p.add_run(confidence_text)
                
                p = doc.add_paragraph()
                p.add_run(f"{t.t('compounds.formula', 'Formula')}: ").bold = True
                p.add_run(formula)
                
                p = doc.add_paragraph()
                p.add_run(f"{t.t('compounds.cas', 'CAS')}: ").bold = True
                p.add_run(cas)
                
                # Añadir imagen 2D
                image_b64 = compound.get('image_2d')
                if image_b64:
                    try:
                        img_bytes = ReportExporter.base64_to_bytes(image_b64)
                        if img_bytes:
                            doc.add_picture(io.BytesIO(img_bytes), width=Inches(1.5))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    except Exception as img_err:
                        print(f"Advertencia: No se pudo renderizar la imagen 2D (DOCX) para {name}: {img_err}")
                
                doc.add_paragraph() # Espacio
            # --- FIN CORRECCIÓN 2 ---
        else:
            doc.add_paragraph(t.t('compounds.none', 'No specific PFAS compounds were detected.'))
        
        # =================================================================
        # === FIN DE LA SECCIÓN MODIFICADA ===
        # =================================================================
        
        doc.add_page_break()

        # PICOS DETECTADOS
        doc.add_heading(f"6. {t.t('report.detected_peaks', 'Detected Peaks')}", level=1)
        
        peaks = results.get("peaks", [])
        
        if peaks:
            peaks_table = doc.add_table(rows=len(peaks) + 1, cols=7)
            peaks_table.style = 'Light Grid Accent 1'
            
            widths = [0.7, 0.8, 0.8, 0.7, 0.7, 0.6, 1.5] # Ajustado ancho
            for i, width in enumerate(widths):
                for cell in peaks_table.columns[i].cells:
                    cell.width = Inches(width)
            
            hdr_cells_peaks = peaks_table.rows[0].cells
            hdr_cells_peaks[0].text = t.t('peaks.ppm', 'PPM')
            hdr_cells_peaks[1].text = t.t('peaks.intensity', 'Intensity')
            hdr_cells_peaks[2].text = t.t('peaks.relative_intensity', 'Rel. Int.')
            hdr_cells_peaks[3].text = t.t('peaks.width', 'Width (ppm)')
            hdr_cells_peaks[4].text = t.t('peaks.width_hz', 'Width (Hz)')
            hdr_cells_peaks[5].text = 'SNR'
            hdr_cells_peaks[6].text = t.t('peaks.region', 'Region')
            
            for idx, peak in enumerate(peaks, 1):
                if idx < len(peaks_table.rows):
                    row_cells = peaks_table.rows[idx].cells
                    
                    ppm_val = peak.get('ppm') or peak.get('position') or 0
                    intensity_val = peak.get('intensity') or peak.get('height') or 0
                    rel_intensity_val = peak.get('relative_intensity') or 0
                    width_ppm_val = peak.get('width_ppm') or peak.get('width') or 0
                    width_hz_val = peak.get('width_hz') or 0
                    snr_val = peak.get('snr') or 0
                    
                    row_cells[0].text = f"{ppm_val:.3f}"
                    row_cells[1].text = f"{intensity_val:,.0f}"
                    row_cells[2].text = f"{rel_intensity_val:.1f}%"
                    row_cells[3].text = f"{width_ppm_val:.3f}"
                    row_cells[4].text = f"{width_hz_val:.1f}"
                    row_cells[5].text = f"{snr_val:.2f}"
                    row_cells[6].text = normalize_region_text(peak.get('region', 'N/A'))
        else:
            doc.add_paragraph(t.t('peaks.none', 'No significant peaks detected.'))

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    # ========================================================================
    # EXPORTACIÓN CSV INDIVIDUAL
    # ========================================================================
    @staticmethod
    def export_csv(results: Dict, lang: str = 'es') -> BinaryIO:
        """Exporta análisis individual como CSV."""
        t = TranslationManager(lang)
        import csv

        text_buffer = io.StringIO()
        writer = csv.writer(text_buffer)

        writer.writerow([t.t('report.title', 'NMR Analysis Report').upper()])
        writer.writerow([t.t('report.sample', 'Sample'), results.get('filename', 'N/A')])
        writer.writerow([t.t('report.date', 'Date'), datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        writer.writerow([])

        # RESUMEN
        writer.writerow([t.t('report.executive_summary', 'Executive Summary').upper()])
        analysis = results.get("analysis", {})
        writer.writerow([t.t('report.parameter', 'Parameter'), t.t('report.value', 'Value'), t.t('report.unit', 'Unit')])
        writer.writerow([t.t('results.fluor', 'Total Fluorine'), f"{analysis.get('fluor_percentage', 0):.2f}", t.t('units.percentage', '%')])
        writer.writerow([t.t('results.pfas', 'PFAS'), f"{analysis.get('pifas_percentage', analysis.get('pfas_percentage', 0)):.2f}", t.t('units.percent_fluor', '% Fluorine')])
        writer.writerow([t.t('results.concentration', 'PFAS Concentration'), f"{analysis.get('pifas_concentration', analysis.get('pfas_concentration', 0)):.4f}", t.t('units.millimolar', 'mM')])
        writer.writerow([t.t('report.quality_score', 'Quality Score'), f"{results.get('quality_score', 0):.1f}", t.t('units.over_ten', '/10')])
        writer.writerow([])

        # PICOS
        peaks = results.get("peaks", [])
        if peaks:
            writer.writerow([t.t('report.detected_peaks', 'Detected Peaks').upper()])
            writer.writerow([t.t('peaks.ppm', 'PPM'), t.t('peaks.intensity', 'Intensity'), t.t('peaks.relative_intensity', 'Rel. Int.'), t.t('peaks.width', 'Width (ppm)'), t.t('peaks.region', 'Region')])
            for peak in peaks:
                width_val = peak.get('width_ppm', peak.get('width', 0))
                writer.writerow([
                     f"{peak.get('ppm', peak.get('position', 0)):.3f}",
                     f"{peak.get('intensity', peak.get('height', 0)):,.0f}",
                     f"{peak.get('relative_intensity', 0):.1f}",
                     f"{width_val:.3f}",
                     normalize_region_text(peak.get('region', 'N/A'))
                ])

        byte_buffer = io.BytesIO()
        byte_buffer.write(text_buffer.getvalue().encode('utf-8'))
        byte_buffer.seek(0)
        text_buffer.close()
        return byte_buffer

    # ========================================================================
    # EXPORTACIÓN DE COMPARACIÓN PDF
    # ========================================================================
    @staticmethod
    def export_comparison_pdf(samples: List[Dict], company_data: Dict, chart_image: bytes = None, lang: str = 'es') -> BinaryIO:
        """Exporta comparación como PDF con gráfico y branding."""
        if not PDF_AVAILABLE: raise ImportError("ReportLab no está instalado.")
        t = TranslationManager(lang)
        output = io.BytesIO()

        doc = BaseDocTemplate(output, pagesize=A4, topMargin=0.75*inch, bottomMargin=1.0*inch, leftMargin=0.75*inch, rightMargin=0.75*inch)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
        template = PageTemplate(id='main', frames=[frame], onPage=lambda canvas, doc: ReportExporter._add_pdf_footer(canvas, doc, company_data))
        doc.addPageTemplates([template])
        story = []
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#2c3e50'), spaceAfter=30, alignment=TA_CENTER)
        subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], alignment=TA_CENTER)
        info_style = ParagraphStyle('Info', parent=styles['Normal'], alignment=TA_CENTER, fontSize=12)

        # PORTADA
        story.append(Paragraph(t.t('comparison.title', "Sample Comparison"), title_style))
        subtitle_text = t.t('comparison.subtitle', f"Comparative Analysis of {len(samples)} Samples")
        if '{count}' in subtitle_text:
            subtitle_text = subtitle_text.replace('{count}', str(len(samples)))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(subtitle_text, subtitle_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"<b>{t.t('report.date', 'Date')}:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", info_style))
        story.append(PageBreak())

        # LOGO
        ReportExporter._add_logo(story, company_data, 'pdf')

        # MUESTRAS INCLUIDAS
        story.append(Paragraph(f"1. {t.t('comparison.samples_included', 'Samples Included')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        sample_list_style = ParagraphStyle('SampleList', parent=styles['Normal'], leftIndent=0.2*inch)
        for idx, sample in enumerate(samples, 1):
            story.append(Paragraph(f"{idx}. {sample.get('filename', sample.get('name', 'N/A'))}", sample_list_style))
        story.append(PageBreak())

        # GRÁFICO COMPARATIVO
        story.append(Paragraph(f"2. {t.t('comparison.chart', 'Comparative Chart')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        if chart_image:
            try:
                img = Image(io.BytesIO(chart_image), width=6.5*inch, height=4*inch)
                img.hAlign = 'CENTER'
                story.append(KeepTogether(img))
            except Exception as e:
                print(f"Error al añadir imagen del gráfico Comp PDF: {e}")
                story.append(Paragraph(f"[{t.t('errors.chartError', 'Error displaying chart')}]", styles['Italic']))
        else:
            story.append(Paragraph(f"[{t.t('errors.noChartData', 'Chart not available')}]", styles['Italic']))
        story.append(PageBreak())

        # TABLA COMPARATIVA
        story.append(Paragraph(f"3. {t.t('comparison.table', 'Comparison Table')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        table_data = [[t.t('report.parameter', 'Parameter')] + [s.get('filename', s.get('name', 'N/A'))[:15] for s in samples]]
        params = [
            (f"{t.t('results.fluor', 'Total Fluorine')} ({t.t('units.percentage', '%')})", 'fluor'),
            (f"{t.t('results.pfas', 'PFAS')} ({t.t('units.percent_fluor', '% Fluorine')})", 'pfas'),
            (f"{t.t('results.concentration', 'PFAS Concentration')} ({t.t('units.millimolar', 'mM')})", 'concentration'),
            (f"{t.t('results.quality', 'Quality')} ({t.t('units.over_ten', '/10')})", 'quality')
        ]
        for param_name, param_key in params:
            row = [param_name]
            for sample in samples:
                value = sample.get(param_key)
                if value is None and param_key == 'pfas': value = sample.get('pifas_percentage')
                if value is None and param_key == 'concentration': value = sample.get('pifas_concentration')
                text_value = '--'
                if value is not None:
                     try:
                         if param_key == 'concentration': text_value = f"{float(value):.4f}"
                         elif param_key == 'quality': text_value = f"{float(value):.1f}"
                         else: text_value = f"{float(value):.2f}"
                     except (ValueError, TypeError): text_value = str(value)
                row.append(text_value)
            table_data.append(row)
        num_cols = len(samples) + 1
        first_col_width = 2.0 * inch
        other_col_width = (doc.width - first_col_width) / len(samples) if samples else 0.5*inch
        col_widths = [first_col_width] + [other_col_width] * len(samples)
        comparison_table = Table(table_data, colWidths=col_widths)
        comparison_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ]))
        story.append(comparison_table)

        doc.build(story)
        output.seek(0)
        return output

    # ========================================================================
    # EXPORTACIÓN DE COMPARACIÓN DOCX
    # ========================================================================
    @staticmethod
    def export_comparison_docx(samples: List[Dict], company_data: Dict, chart_image: bytes = None, lang: str = 'es') -> BinaryIO:
        """Exporta comparación como DOCX con gráfico y branding."""
        if not DOCX_AVAILABLE: raise ImportError("python-docx no está instalado.")
        t = TranslationManager(lang)
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # PIE DE PÁGINA
        ReportExporter._add_docx_footer(doc, company_data)

        # PORTADA
        title = doc.add_heading(t.t('comparison.title', "Sample Comparison"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_text = t.t('comparison.subtitle', f"Comparative Analysis of {len(samples)} Samples")
        if '{count}' in subtitle_text:
            subtitle_text = subtitle_text.replace('{count}', str(len(samples)))
        subtitle = doc.add_heading(subtitle_text, level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        date_p = doc.add_paragraph(f"{t.t('report.date', 'Date')}: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # LOGO
        ReportExporter._add_logo(doc, company_data, 'docx')

        # MUESTRAS INCLUIDAS
        doc.add_heading(f"1. {t.t('comparison.samples_included', 'Samples Included')}", level=1)
        for idx, sample in enumerate(samples, 1):
            try:
                doc.add_paragraph(f"{idx}. {sample.get('filename', sample.get('name', 'N/A'))}", style='List Number')
            except KeyError:
                doc.add_paragraph(f"{idx}. {sample.get('filename', sample.get('name', 'N/A'))}")
        doc.add_page_break()

        # GRÁFICO COMPARATIVO
        doc.add_heading(f"2. {t.t('comparison.chart', 'Comparative Chart')}", level=1)
        if chart_image:
            try:
                doc.add_picture(io.BytesIO(chart_image), width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error al añadir imagen del gráfico Comp DOCX: {e}")
                doc.add_paragraph(f"[{t.t('errors.chartError', 'Error displaying chart')}]")
        else:
            doc.add_paragraph(f"[{t.t('errors.noChartData', 'Chart not available')}]")
        doc.add_page_break()

        # TABLA COMPARATIVA
        doc.add_heading(f"3. {t.t('comparison.table', 'Comparison Table')}", level=1)
        num_params = 4
        table = doc.add_table(rows=num_params + 1, cols=len(samples) + 1)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = t.t('report.parameter', 'Parameter')
        for idx, sample in enumerate(samples, 1):
            hdr_cells[idx].text = sample.get('filename', sample.get('name', 'N/A'))[:20]
        params = [
            (f"{t.t('results.fluor', 'Total Fluorine')} ({t.t('units.percentage', '%')})", 'fluor'),
            (f"{t.t('results.pfas', 'PFAS')} ({t.t('units.percent_fluor', '% Fluorine')})", 'pfas'),
            (f"{t.t('results.concentration', 'PFAS Concentration')} ({t.t('units.millimolar', 'mM')})", 'concentration'),
            (f"{t.t('results.quality', 'Quality')} ({t.t('units.over_ten', '/10')})", 'quality')
        ]
        for row_idx, (param_name, param_key) in enumerate(params, 1):
             if row_idx < len(table.rows):
                 row_cells = table.rows[row_idx].cells
                 row_cells[0].text = param_name
                 for col_idx, sample in enumerate(samples, 1):
                     value = sample.get(param_key)
                     if value is None and param_key == 'pfas': value = sample.get('pifas_percentage')
                     if value is None and param_key == 'concentration': value = sample.get('pifas_concentration')
                     text_value = '--'
                     if value is not None:
                         try:
                             if param_key == 'concentration': text_value = f"{float(value):.4f}"
                             elif param_key == 'quality': text_value = f"{float(value):.1f}"
                             else: text_value = f"{float(value):.2f}"
                         except (ValueError, TypeError): text_value = str(value)
                     row_cells[col_idx].text = text_value

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    # ========================================================================
    # EXPORTACIÓN DE COMPARACIÓN CSV
    # ========================================================================
    @staticmethod
    def export_comparison_csv(samples: List[Dict], lang: str = 'es') -> BinaryIO:
         """Exporta comparación como CSV."""
         t = TranslationManager(lang)
         import csv
         text_buffer = io.StringIO()
         writer = csv.writer(text_buffer)
         writer.writerow([t.t('comparison.title', "Sample Comparison").upper()])
         writer.writerow([t.t('report.date', 'Date'), datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
         writer.writerow([t.t('comparison.samples_included', 'Samples Included'), len(samples)])
         writer.writerow([])
         headers = [t.t('report.parameter', 'Parameter')] + [s.get('filename', s.get('name', 'N/A')) for s in samples]
         writer.writerow(headers)
         params = [
            (f"{t.t('results.fluor', 'Total Fluorine')} ({t.t('units.percentage', '%')})", 'fluor'),
            (f"{t.t('results.pfas', 'PFAS')} ({t.t('units.percent_fluor', '% Fluorine')})", 'pfas'),
            (f"{t.t('results.concentration', 'PFAS Concentration')} ({t.t('units.millimolar', 'mM')})", 'concentration'),
            (f"{t.t('results.quality', 'Quality')} ({t.t('units.over_ten', '/10')})", 'quality')
         ]
         for param_name, param_key in params:
             row = [param_name]
             for sample in samples:
                 value = sample.get(param_key)
                 if value is None and param_key == 'pfas': value = sample.get('pifas_percentage')
                 if value is None and param_key == 'concentration': value = sample.get('pifas_concentration')
                 text_value = '--'
                 if value is not None:
                     try:
                         if param_key == 'concentration': text_value = f"{float(value):.4f}"
                         elif param_key == 'quality': text_value = f"{float(value):.1f}"
                         else: text_value = f"{float(value):.2f}"
                     except (ValueError, TypeError): text_value = str(value)
                 row.append(text_value)
             writer.writerow(row)
         byte_buffer = io.BytesIO()
         byte_buffer.write(text_buffer.getvalue().encode('utf-8'))
         byte_buffer.seek(0)
         text_buffer.close()
         return byte_buffer

    # ========================================================================
    # EXPORTACIÓN DE DASHBOARD PDF
    # ========================================================================
    @staticmethod
    def export_dashboard_pdf(stats: Dict, company_data: Dict, chart_images: Dict[str, bytes] = None, lang: str = 'es') -> BinaryIO:
        """Exporta dashboard como PDF con gráficos y branding."""
        if not PDF_AVAILABLE: raise ImportError("ReportLab no está instalado.")
        t = TranslationManager(lang)
        output = io.BytesIO()

        doc = BaseDocTemplate(output, pagesize=A4, topMargin=0.75*inch, bottomMargin=1.0*inch, leftMargin=0.75*inch, rightMargin=0.75*inch)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
        template = PageTemplate(id='main', frames=[frame], onPage=lambda canvas, doc: ReportExporter._add_pdf_footer(canvas, doc, company_data))
        doc.addPageTemplates([template])
        story = []
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#2c3e50'), spaceAfter=30, alignment=TA_CENTER)
        subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], alignment=TA_CENTER)
        info_style = ParagraphStyle('Info', parent=styles['Normal'], alignment=TA_CENTER, fontSize=12)

        # PORTADA
        story.append(Paragraph(t.t('dashboard.title', "Dashboard"), title_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(t.t('dashboard.subtitle', "Statistical Summary"), subtitle_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"<b>{t.t('report.date', 'Date')}:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", info_style))
        story.append(PageBreak())

        # LOGO
        ReportExporter._add_logo(story, company_data, 'pdf')

        # ESTADÍSTICAS GENERALES
        story.append(Paragraph(f"1. {t.t('dashboard.general_stats', 'General Statistics')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        stats_report_data = [
            [t.t('report.metric', 'Metric'), t.t('report.value', 'Value')],
            [t.t('dashboard.totalAnalyses', 'Total Analyses'), str(stats.get('totalAnalyses', 0))],
            [t.t('dashboard.avgFluor', 'Average Fluorine'), f"{stats.get('avgFluor', 0):.2f}{t.t('units.percentage', '%')}"],
            [t.t('dashboard.avgPfas', 'Average PFAS'), f"{stats.get('avgPfas', 0):.2f}{t.t('units.percentage', '%')}"],
            [t.t('dashboard.avgConcentration', 'Average Concentration'), f"{stats.get('avgConcentration', 0):.4f} {t.t('units.millimolar', 'mM')}"],
            [t.t('dashboard.avgQuality', 'Average Quality'), f"{stats.get('avgQuality', 0):.1f}{t.t('units.over_ten', '/10')}"],
            [t.t('dashboard.highQuality', 'High Quality Samples'), str(stats.get('highQualitySamples', 0))],
            [t.t('dashboard.avg_snr', 'Average SNR'), f"{stats.get('avgSNR', stats.get('avg_snr', 0)):.2f}" if stats.get('avgSNR', stats.get('avg_snr')) is not None else 'N/A']
        ]
        stats_table = Table(stats_report_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(stats_table)
        story.append(PageBreak())

        # GRÁFICOS
        story.append(Paragraph(f"2. {t.t('dashboard.charts', 'Charts')}", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        if chart_images:
            if 'trend' in chart_images and chart_images['trend']:
                 story.append(Paragraph(f"2.1 {t.t('dashboard.trend', 'Trend')}", styles['Heading2']))
                 story.append(Spacer(1, 0.1*inch))
                 try:
                     img = Image(io.BytesIO(chart_images['trend']), width=6*inch, height=3.5*inch)
                     img.hAlign='CENTER'
                     story.append(KeepTogether(img))
                 except Exception as e:
                     print(f"Error al añadir gráfico Trend PDF: {e}")
                 story.append(Spacer(1, 0.3*inch))

            if 'distribution' in chart_images and chart_images['distribution']:
                 story.append(Paragraph(f"2.2 {t.t('dashboard.distribution', 'Distribution')}", styles['Heading2']))
                 story.append(Spacer(1, 0.1*inch))
                 try:
                     img = Image(io.BytesIO(chart_images['distribution']), width=6*inch, height=3.5*inch)
                     img.hAlign='CENTER'
                     story.append(KeepTogether(img))
                 except Exception as e:
                     print(f"Error al añadir gráfico Dist PDF: {e}")
        else:
            story.append(Paragraph(f"[{t.t('errors.noChartData', 'Charts not available')}]", styles['Italic']))

        doc.build(story)
        output.seek(0)
        return output

    # ========================================================================
    # EXPORTACIÓN DE DASHBOARD DOCX
    # ========================================================================
    @staticmethod
    def export_dashboard_docx(stats: Dict, company_data: Dict, chart_images: Dict[str, bytes] = None, lang: str = 'es') -> BinaryIO:
        """Exporta dashboard como DOCX con gráficos y branding."""
        if not DOCX_AVAILABLE: raise ImportError("python-docx no está instalado.")
        t = TranslationManager(lang)
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # PIE DE PÁGINA
        ReportExporter._add_docx_footer(doc, company_data)

        # PORTADA
        title = doc.add_heading(t.t('dashboard.title', "Dashboard"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading(t.t('dashboard.subtitle', "Statistical Summary"), level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        date_p = doc.add_paragraph(f"{t.t('report.date', 'Date')}: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # LOGO
        ReportExporter._add_logo(doc, company_data, 'docx')

        # ESTADÍSTICAS GENERALES
        doc.add_heading(f"1. {t.t('dashboard.general_stats', 'General Statistics')}", level=1)
        stats_report_data = [
            (t.t('dashboard.totalAnalyses', 'Total Analyses'), str(stats.get('totalAnalyses', 0))),
            (t.t('dashboard.avgFluor', 'Average Fluorine'), f"{stats.get('avgFluor', 0):.2f}{t.t('units.percentage', '%')}" ),
            (t.t('dashboard.avgPfas', 'Average PFAS'), f"{stats.get('avgPfas', 0):.2f}{t.t('units.percentage', '%')}"),
            (t.t('dashboard.avgConcentration', 'Average Concentration'), f"{stats.get('avgConcentration', 0):.4f} {t.t('units.millimolar', 'mM')}"),
            (t.t('dashboard.avgQuality', 'Average Quality'), f"{stats.get('avgQuality', 0):.1f}{t.t('units.over_ten', '/10')}"),
            (t.t('dashboard.highQuality', 'High Quality Samples'), str(stats.get('highQualitySamples', 0))),
            (t.t('dashboard.avg_snr', 'Average SNR'), f"{stats.get('avgSNR', stats.get('avg_snr', 0)):.2f}" if stats.get('avgSNR', stats.get('avg_snr')) is not None else 'N/A')
        ]
        table = doc.add_table(rows=len(stats_report_data) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells_dash = table.rows[0].cells
        hdr_cells_dash[0].text = t.t('report.metric', 'Metric')
        hdr_cells_dash[1].text = t.t('report.value', 'Value')
        for idx, (metric, value) in enumerate(stats_report_data, 1):
             if idx < len(table.rows):
                 row_cells = table.rows[idx].cells
                 row_cells[0].text = metric
                 row_cells[1].text = value
        doc.add_page_break()

        # GRÁFICOS
        doc.add_heading(f"2. {t.t('dashboard.charts', 'Charts')}", level=1)
        if chart_images:
             if 'trend' in chart_images and chart_images['trend']:
                 doc.add_heading(f"2.1 {t.t('dashboard.trend', 'Trend')}", level=2)
                 try:
                     doc.add_picture(io.BytesIO(chart_images['trend']), width=Inches(6))
                     doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                     doc.add_paragraph()
                 except Exception as e:
                     print(f"Error al añadir gráfico Trend DOCX: {e}")

             if 'distribution' in chart_images and chart_images['distribution']:
                 doc.add_heading(f"2.2 {t.t('dashboard.distribution', 'Distribution')}", level=2)
                 try:
                     doc.add_picture(io.BytesIO(chart_images['distribution']), width=Inches(6))
                     doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                 except Exception as e:
                     print(f"Error al añadir gráfico Dist DOCX: {e}")
        else:
            doc.add_paragraph(f"[{t.t('errors.noChartData', 'Charts not available')}]")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output